# -*- coding: utf-8 -*-
"""Tạo báo cáo CHƯƠNG 5 (chuong5.docx). Chạy: python scripts/generate_chuong5_docx.py"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "chuong5.docx"


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    for t in items:
        p = doc.add_paragraph(t, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        for r in p.runs:
            r.font.size = Pt(12)


def main():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Pt(72)

    t = doc.add_heading("CHƯƠNG 5. CÀI ĐẶT VÀ TRIỂN KHAI", level=1)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # --- 5.1 ---
    doc.add_heading("5.1. Môi trường và triển khai", level=2)

    add_para(
        doc,
        "Dự án LMS được tổ chức thành hai thành phần chính trong cùng một kho mã: "
        "thư mục doan_be-main (backend API) và doan_fe-main (ứng dụng web SPA). "
        "Triển khai thử nghiệm phát triển (development) mặc định chạy trực tiếp trên Node.js "
        "và trình duyệt — không bắt buộc container; có thể bổ sung Docker/Kubernetes trong "
        "ghi chú production nếu yêu cầu nhân rộng.",
    )

    add_para(doc, "Phần cứng / máy chủ gợi ý:", bold=True)
    add_bullets(
        doc,
        [
            "Môi trường dev: máy trạm hoặc laptop Windows/Linux/macOS, RAM ≥ 8 GB để đồng thời chạy PostgreSQL và hai tiến trình Node (FE + BE).",
            "Môi trường pilot/production: VPS hoặc cloud VM (Ubuntu LTS,) tối thiểu 2 vCPU và 4 GB RAM cho một số ít học viên đồng thời; tách riêng DB hoặc dùng dịch vụ RDS tương đương.",
            "Đĩa: đủ dung lượng cho WAL PostgreSQL và file media MinIO khi kích hoạt lưu trữ đối tượng.",
        ],
    )

    add_para(doc, "Phiên bản runtime và công nghệ lõi:", bold=True)
    add_bullets(
        doc,
        [
            "Node.js 20+ (LTS được khuyến nghị; Vite 6 tương thích với các bản Node hiện dùng của dự án).",
            "PostgreSQL 14+ làm DB chính; client `pg` (driver) trên backend.",
            "Backend: Express 5, TypeScript, `ts-node`/`nodemon` khi dev, script migration/seed qua `tsx`.",
            "Frontend: React 19, Vite, TypeScript, Ant Design 6, Tailwind 4, axios, React Router 7.",
            "MinIO (tùy chọn): cổng API lưu trữ thường 9000 — tránh nhầm với PostgreSQL.",
        ],
    )

    add_para(doc, "Container (Docker) — không bắt buộc:", bold=True)
    add_para(
        doc,
        "Dự án có thể triển khai không có Docker: cài Postgres và Node trực tiếp. "
        "Nếu chuẩn hóa bằng container, có thể đặt Dockerfile cho backend (multi-stage build `npm ci && npm run build` nếu sau này dùng bản transpile JS), "
        "frontend phục vụ bằng `nginx` tĩnh sau `vite build`, và `compose` ghép với Postgres + MinIO. "
        "Nội dung báo cáo này mô tả triển khai trực tiếp làm đường cơ sở.",
    )

    add_para(doc, "Biến môi trường backend (doan_be-main):", bold=True)
    add_bullets(
        doc,
        [
            "PORT: cổng HTTP (ví dụ 8096); API gốc thường là http://host:PORT/api.",
            "DB_HOST, DB_PORT, DB_NAME, DB_USERNAME (hoặc DB_USER), DB_PASSWORD: kết nối PostgreSQL — script migrate/seed đọc từ .env trong thư mục backend.",
            "JWT_SECRET, JWT_EXPIRES_IN, các biến refresh token theo cấu hình `config.ts/jwt.ts` của dự án.",
            "MinIO: MINIO_ENDPOINT, MINIO_PORT, MINIO_ACCESS_KEY, MINIO_SECRET_KEY, MINIO_BUCKET, … có thể tắt bằng MINIO_DISABLED=true khi chỉ kiểm thử không upload.",
            "SEED_DEMO_PASSWORD, SEED_PURGE (tùy): cho script seed-demo khi nhập liệu mẫu.",
        ],
    )

    add_para(doc, "Biến môi trường frontend (doan_fe-main):", bold=True)
    add_bullets(
        doc,
        [
            "Thường có biến trỏ base URL của API (ví dụ VITE_API_URL hoặc tương đương trong .env) để SPA gọi backend khi không dùng proxy Vite.",
            "Chế độ dev: npm run dev (Vite, mặc định cổng 5173).",
        ],
    )

    add_para(doc, "Lưu ý thống nhất đường dẫn .env:", bold=True)
    add_para(
        doc,
        "Một số file khởi động server có thể đọc dotenv từ thư mục cha của backend; cần đảm bảo file .env "
        "khớp đường dẫn khi chạy `npm start` và khi migrate/seed — tránh sai lệch biến DB giữa runtime và tooling.",
        italic=False,
    )

    add_para(doc, "Luồng cài đặt cơ bản (ghi nhớ cho triển khai):", bold=True)
    add_bullets(
        doc,
        [
            "git clone hoặc sao chép mã → cd doan_be-main && npm install.",
            "Cấu hình .env → npm run db:migrate → npm run db:seed (tùy) → npm start.",
            "cd doan_fe-main && npm install → npm run dev (dev) hoặc npm run build + phục vụ tệp dist (production preview).",
            "Điều chỉnh CORS trong app backend cho đúng origin FE sau khi deploy domain thật.",
        ],
    )

    # --- 5.2 ---
    doc.add_heading("5.2. Triển khai chức năng lõi LMS", level=2)

    add_para(
        doc,
        "Chức năng lõi đã gắn với lớp router Express, controller/service/repository (DI qua tsyringe) "
        "và schema PostgreSQL qua migration script theo thứ tự tên file.",
    )

    add_para(doc, "Quản lý khóa học và nội dung:", bold=True)
    add_bullets(
        doc,
        [
            "Tạo/sửa/xóa/xuất bản khóa học và upload thumbnail/video (khi MinIO khả dụng).",
            "Đồng bộ curriculum: sections và lessons — hỗ trợ các loại bài lesson_type video, docs, quiz, exam với content_json (JSONB).",
            "Tra cứu catalog công khai: endpoint POST `/course/catalog-published` (không middleware auth), dùng hàm SQL `fn_list_published_courses` (migration V022) để phân trang, lọc category/major, tìm kiếm.",
            "Đăng ký học và (khi có) thanh toán/mua khóa: enrollment với điều kiện trạng thái published và quyền truy cập full/preview.",
            "Learning paths: bundle nhiều khóa theo learning_path_items (migration V019/V020/V021 và seed đi kèm).",
        ],
    )

    add_para(doc, "Xác thực và phân quyền:", bold=True)
    add_bullets(
        doc,
        [
            "Xác thực HTTP-only cookie + JWT (access/refresh) theo triển khai hiện tại; axios FE gửi `withCredentials`.",
            "authMiddleware và optionalAuthMiddleware: bảo vệ route cần đăng nhập hoặc bổ sung ngữ cảnh user.",
            "RBAC: bảng roles, user_roles và (khi triển khai đầy đủ) permissions/role_permissions — phân tách vai trò admin / instructor / student.",
            "Instructor chỉnh sửa nội dung của chính mình trong phạm vi nghiệp vụ được kiểm soát trong service layer.",
            "Khu `/admin`: layout và route được bảo vệ theo role admin (routing React + kiểm tra API phía server).",
        ],
    )

    # --- 5.3 ---
    doc.add_heading("5.3. Triển khai thu thập tiến độ và gợi ý AI", level=2)

    add_para(
        doc,
        "Tầng dữ liệu hành vi đã có bảng lesson_progress, user_quiz_attempts và study_events ( migration V006). "
        "Việc triển khai thu gom đầy đủ và module gợi ý có thể chia thành giai đoạn: "
        "(1) ghi nhận tiến độ & kết quả bắt buộc cho LMS, (2) bổ sung sự kiện chi tiết, (3) pipeline tính toán và dịch vụ gợi ý.",
    )

    add_para(doc, "Ghi nhận sự kiện và tiến độ (triển khai thực tế / mở rộng):", bold=True)
    add_bullets(
        doc,
        [
            "lesson_progress: upsert khi người học xem nội dung / hoàn thành bài — nguồn chân lý tiến độ theo lesson.",
            "course_enrollments.progress_percent và completed_at được cập nhật đồng bộ khi có quy tắc rollup từ lesson → course.",
            "user_quiz_attempts: mỗi lần nộp quiz/exam tạo bản ghi với điểm, passed và answers JSONB để tái hiện/phân tích.",
            "study_events: gọi có chủ đích từ client hoặc từ server hook (lesson_open, video_heartbeat…) với event_type và duration_seconds trong payload có schema rõ.",
        ],
    )

    add_para(doc, "Pipeline tính toán chỉ báo định kỳ hoặc theo yêu cầu:", bold=True)
    add_bullets(
        doc,
        [
            "Định kỳ (batch nightly): SQL hoặc job Node aggregate theo user_id/tuần: tổng thời lượng, số ngày hoạt động, điểm trung bình quiz theo chủ đề.",
            "Theo yêu cầu (on-demand): API internal sau mỗi quiz_submit tái tính vector chủ đề hoặc cache gợi ý TTL ngắn.",
            "Có thể dùng materialized view user_learning_summary(user_id, period, features jsonb) và REFRESH theo lịch.",
            "Chuẩn hóa timezone UTC lưu DB; dashboard hiển thị theo múi giờ người dùng ở FE.",
        ],
    )

    add_para(doc, "Tích hợp dịch vụ gợi ý với LMS:", bold=True)
    add_bullets(
        doc,
        [
            "Tách module Recommendation thành lớp service (TypeScript) hoặc microservice nhỏ: đầu vào profile + chỉ báo (Chương 4), đầu ra danh sách khóa/lộ trình có rationale.",
            "Điểm tích hợp trong FE: một route/component “Đề xuất” gọi GET/POST `/recommendations` (định nghĩa khi triển khai), hiển thị cạnh catalog và “Tiếp tục học”.",
            "Fallback khi chưa có model: kết hợp phổ biến (enrollment_count từ catalog SQL), lọc category/major đã tương tác, và bước kế tiếp trong learning_paths.",
            "Khi nâng cấp ML: triển khai training pipeline ngoài request path; phiên bản model và feature store phiên bản hóa để rollback.",
        ],
    )

    # --- 5.4 ---
    doc.add_heading("5.4. Dashboard", level=2)

    add_para(
        doc,
        "Dashboard người dùng và dashboard quản trị kết nối với cùng nguồn dữ liệu backend (PostgreSQL truy vấn qua API). "
        "Frontend đã dùng Recharts / Ant Design cho biểu đồ và bảng — dưới đây là cách nối nguồn dữ liệu với phác thảo Chương 4.",
    )

    add_para(doc, "Dashboard người học (/dashboard và các khối liên quan):", bold=True)
    add_bullets(
        doc,
        [
            "Nguồn: các API bọc enrollment, tiến độ tổng hợp, lịch sử làm quiz (tổng hợp từ lesson_progress và user_quiz_attempts).",
            "Widget “tiếp tục học”: lấy danh sách enrollment + last_access nếu bổ sung trường hoặc suy luận từ lesson_progress.updated_at.",
            "Widget chủ đề: join course → course_categories để nhóm chỉ báo theo nhóm khóa hướng nghiệp.",
            "Gợi ý AI: nút/card gọi endpoint recommendation và cache client ngắn hạn để giảm tải.",
        ],
    )

    add_para(doc, "Dashboard quản trị (/admin/dashboard):", bold=True)
    add_bullets(
        doc,
        [
            "Chỉ user có role admin; API trả về tổng số khóa, user, phiên làm việc học trong khoảng thời gian (study_events aggregated khi có dữ liệu).",
            "Biểu đồ xu hướng đăng ký: query course_enrollments group by DATE(enrolled_at).",
            "Báo cáo “sức khỏe nội dung”: quiz pass rate theo lesson_id để cờ bài khó quá/thiếu tương tác.",
        ],
    )

    add_para(doc, "Đồng bộ và an toàn khi báo cáo:", bold=True)
    add_bullets(
        doc,
        [
            "Phân quyền kiểm tra hai phía — không chỉ ẩn menu FE mà middleware backend chặn dữ liệu nhạy cảm aggregat cho role sai.",
            "Đối với chỉ báo thô, hạn chế export PII; log truy vấn chậm để tuning index (idx đã có trên study_events user/time và lesson_progress user).",
        ],
    )

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(
        "Ghi chú: chương này mô tả trình tự và phương án triển khai gắn với codebase LMS hiện có (doan_be-main, doan_fe-main). "
        "Các đoạn pipeline gợi ý và API dashboard mở rộng được trình bày như bước triển khai tiếp theo có thể tách sprint; "
        "phần hạ tầng container có thể bổ sung khi có yêu cầu SLA và CI/CD chính thức."
    ).italic = True

    doc.save(OUT)
    print("Wrote:", OUT)


if __name__ == "__main__":
    main()
