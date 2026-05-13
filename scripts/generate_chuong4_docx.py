# -*- coding: utf-8 -*-
"""Tạo báo cáo CHƯƠNG 4 (chuong4.docx). Chạy: python scripts/generate_chuong4_docx.py"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "chuong4.docx"


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    for t in items:
        p = doc.add_paragraph(t, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        for r in p.runs:
            r.font.size = Pt(12)


def main():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Pt(72)

    t = doc.add_heading("CHƯƠNG 4. THIẾT KẾ HỆ THỐNG", level=1)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading("4.1. Kiến trúc tổng thể", level=2)
    body = doc.add_paragraph()
    body.add_run(
        "Hệ thống LMS hướng nghiệp được tổ chức theo mô hình "
    )
    body.add_run("Client–Server").bold = True
    body.add_run(
        " (tách rõ tầng giao diện và tầng dịch vụ), trong đó tầng giao diện "
        "(React/Vite/Ant Design) chỉ chứa trình duyệt và trạng thái hiển thị; "
        "tầng máy chủ (Node.js/Express, PostgreSQL, lưu trữ media qua MinIO khi có) "
        "chịu trách nhiệm nghiệp vụ, bảo vệ đường truyền API và nhất quán dữ liệu."
    )
    body.paragraph_format.space_after = Pt(8)

    add_para(doc, "Luồng dữ liệu phục vụ học tập và gợi ý (thiết kế mục tiêu):")

    add_bullets(
        doc,
        [
            "Ghi nhận học: người dùng tương tác bài học (tiến độ, video, nội dung docs, quiz/exam) → "
            "API cập nhật lesson_progress, user_quiz_attempts và (khi bật thu thập chi tiết) study_events "
            "với payload chuẩn hóa.",
            "Kho chỉ báo / feature: job hoặc truy vấn tổng hợp định kỳ (hoặc theo sự kiện) tính các đại lượng "
            "theo user_id: tỷ lệ hoàn thành khóa, điểm trung bình trắc nghiệm, thời lượng học ước lượng, "
            "tần suất đăng nhập học, chủ đề (category/major/tag) được tương tác.",
            "Động cơ gợi ý: module AI Recommendation đọc feature vector/profile người học + luật/heuristic/content-based "
            "hoặc mô hình ML (tùy giai đoạn triển khai), xếp hạng khóa học/lộ trình phù hợp và giải thích ngắn.",
            "Phản hồi API/UI: kết quả gợi ý được trả qua các endpoint và hiển thị tại các màn hình "
            "“Gợi ý/Bạn có thể quan tâm”, “Lộ trình cá nhân”, và bổ sung thẻ contextual trên catalog.",
        ],
    )

    add_para(doc, "")
    fig = doc.add_paragraph()
    fig.add_run("Sơ đồ luồng logic (mức khối)\n").bold = True
    fig.add_run(
        "[Client] ──HTTPS──▶ [API Gateway / Express Routes + Auth Middleware]\n"
        "                              │\n"
        "                              ├──▶ [PostgreSQL: users/roles/courses/… và bảng hành vi học]\n"
        "                              ├──▶ [MinIO/Media (video, thumbnail…)]\n"
        "                              └──▶ [Recommendation Service] ◀── feature store / aggregated views\n"
    )
    fig.paragraph_format.space_after = Pt(8)

    # --- 4.2 ---
    doc.add_heading("4.2. Thiết kế cơ sở dữ liệu và ERD", level=2)

    add_para(doc, "Nguyên tắc thiết kế:")
    add_bullets(
        doc,
        [
            "Phân tách thực thể “nội dung” (khóa học, chương, bài) và thực thể “hành vi” (đăng ký, tiến độ, làm bài, sự kiện học).",
            "Hỗ trợ gợi ý thông qua đa liên kết chủ đề: major, category, learning path (bundle), reviews/wishlist (khi khai thác).",
            "Cho phép mở rộng log hành vi (JSONB) mà không phá vỡ schema cốt lõi.",
        ],
    )

    add_para(doc, "Nhóm bảng chính (theo migration thực tế của dự án):")

    add_para(doc, "Người dùng và phân quyền", bold=True)
    add_bullets(
        doc,
        [
            "users: định danh, thông tin hồ sơ, trạng thái tài khoản.",
            "roles, permissions, role_permissions, user_roles: RBAC; phân biệt admin / instructor / student.",
            "sessions: phiên đăng nhập, refresh token (an toàn truy cập).",
        ],
    )

    add_para(doc, "Khóa học và cấu trúc nội dung", bold=True)
    add_bullets(
        doc,
        [
            "courses, sections, lessons: cây khóa học; lesson_type (video, docs, quiz, exam); content_json cho nội dung/bộ câu hỏi.",
            "major, categories: chủ đề/ngành/danh mục “tag” tĩnh.",
            "course_majors, course_categories: quan hệ nhiều–nhiều — đóng vai trò “tag chủ đề” phục vụ lọc catalog và gợi ý content-based.",
            "course_enrollments: quan hệ user–course, progress_percent, completed_at, access_kind (theo mở rộng LMS).",
        ],
    )

    add_para(doc, "Tiến độ, kết quả kiểm tra và sự kiện học", bold=True)
    add_bullets(
        doc,
        [
            "lesson_progress: tiến độ theo bài, is_completed — nền tảng tính completion rate.",
            "user_quiz_attempts: nhiều lần làm quiz/exam; score, max_score, passed, answers (JSONB), mốc thời gian.",
            "study_events: log sự kiện (event_type, duration_seconds, payload JSONB) phục vụ phân tích hành vi chi tiết.",
        ],
    )

    add_para(doc, "Lộ trình, thương mại, chứng chỉ, đánh giá (mở rộng)", bold=True)
    add_bullets(
        doc,
        [
            "learning_paths, learning_path_items: combo khóa theo thứ tự — “lộ trình” gợi ý có cấu trúc.",
            "path_purchases, course_purchases: tín hiệu giá trị kinh tế / cam kết học.",
            "certificates, course_reviews: kết quả đầu ra và phản hồi chất lượng.",
        ],
    )

    add_para(doc, "ERD (mô tả quan hệ)", bold=True)
    erd = doc.add_paragraph()
    erd.add_run(
        "users 1─N user_roles N─1 roles; users 1─N course_enrollments N─1 courses; "
        "courses 1─N sections 1─N lessons; users 1─N lesson_progress N─1 lessons; "
        "users 1─N user_quiz_attempts N─1 lessons; users 1─N study_events; "
        "courses N─N categories qua course_categories; courses N─N major qua course_majors; "
        "learning_paths 1─N learning_path_items N─1 courses. "
        "Chi tiết khóa ngoại và ràng buộc CASCADE được khai báo trong các file migration PostgreSQL của dự án."
    )
    erd.paragraph_format.space_after = Pt(8)

    doc.add_heading("4.3. Chỉ báo và thu thập dữ liệu phục vụ phân tích", level=2)

    add_para(
        doc,
        "Mục tiêu: biến dữ liệu thô học tập thành vector chỉ báo ổn định, có thể so sánh giữa người học và phục vụ dashboard + gợi ý.",
    )

    add_para(doc, "Định nghĩa chỉ báo (ví minh họa đặt tên I1, I2…):", bold=True)

    tbl = doc.add_table(rows=1, cols=3)
    hdr = tbl.rows[0].cells
    hdr[0].text = "Mã"
    hdr[1].text = "Ý nghĩa / công thức gợi ý"
    hdr[2].text = "Nguồn dữ liệu"
    rows = [
        (
            "I1",
            "Kết quả học–kiểm tra: điểm quiz/exam chuẩn hóa theo max_score; tỉ lệ passed; điểm trung bình theo chủ đề (category_id).",
            "user_quiz_attempts (+ join lessons → course → course_categories)",
        ),
        (
            "I2",
            "Mức hoàn thành khóa: progress_percent tại course_enrollments; hoặc Σ(is_completed lesson) / COUNT(lesson trong khóa).",
            "course_enrollments, lesson_progress, sections/lessons",
        ),
        (
            "I3",
            "Thời lượng học (ước lượng): tổng duration_seconds theo user/khóa/bài; hoặc chênh lệch thời gian session nếu triển khai heartbeat.",
            "study_events (event_type, duration_seconds), hoặc video progress nếu bổ sung",
        ),
        (
            "I4",
            "Tần suất học: số ngày có hoạt động trong cửa sổ 7/30 ngày; streak học liên tiếp.",
            "study_events.created_at hoặc aggregat từ lesson_progress.updated_at",
        ),
        (
            "I5",
            "Độ tương tác chủ đề: vector trọng số theo category_id/major_id dựa trên thời gian hoặc số bài đã hoàn thành.",
            "lesson_progress + course_categories + course_majors",
        ),
        (
            "I6",
            "Tín hiệu cam kết: đã mua path/course (path_purchases, course_purchases), wishlist nếu có.",
            "course_purchases, path_purchases, course_wishlists (nếu bật)",
        ),
    ]
    for m, y, n in rows:
        r = tbl.add_row().cells
        r[0].text = m
        r[1].text = y
        r[2].text = n

    add_para(doc, "")
    add_para(doc, "Chiến lược ghi log và chuẩn hóa", bold=True)
    add_bullets(
        doc,
        [
            "Mức 1 (bắt buộc): mỗi lần cập nhật tiến độ bài/khóa ghi lesson_progress và sync course_enrollments.progress_percent theo quy tắc nghiệp vụ.",
            "Mức 2 (khuyến nghị): study_events với event_type thuộc tập hợp cố định (vd. lesson_open, lesson_close, video_progress, quiz_submit) "
            "và payload JSON schema versioned (vd. schemaVersion: 1).",
            "Chuẩn hóa thời gian: lưu TIMESTAMPTZ UTC; xử lý báo cáo theo múi giờ ở tầng API.",
            "Ẩn danh hóa / tối thiểu hóa PII trong payload log; không ghi mật khẩu hay token.",
            "Batch aggregation: ETL định kỳ (hoặc materialized view) sinh bảng user_learning_summary(user_id, window, jsonb features) để giảm tải truy vấn dashboard.",
        ],
    )

    doc.add_heading("4.4. Thiết kế module AI Recommendation", level=2)

    add_para(
        doc,
        "Module được thiết kế theo hướng modul hóa: có thể khởi động bằng heuristic/content-based, "
        "sau đó nâng cấp embedding hoặc mô hình học có giám sát khi đủ dữ liệu.",
    )

    add_para(doc, "Đầu vào (Input)", bold=True)
    add_bullets(
        doc,
        [
            "User profile: role, optional khối/lớp hoặc chuyên ngành đăng ký (nếu bổ sung trên users hoặc bảng mở rộng).",
            "Implicit feedback: vectơ chỉ báo I1–I6, tiến độ theo category/major, lịch sử enroll.",
            "Explicit feedback (tùy chọn): đánh giá khóa, wishlist.",
            "Item metadata: embedding one-hot/category tags cho courses; textual title+description để học semantic model sau này.",
            "Ngữ cảnh thời gian: học kỳ, mục tiêu (“ôn THPT”, “tìm việc”) nếu thu thập qua onboarding.",
        ],
    )

    add_para(doc, "Thuật toán / heuristic (giai đoạn 1 – đề xuất triển khai)", bold=True)
    add_bullets(
        doc,
        [
            "Content-based scoring: similarity(user_topic_vector, course_topic_vector) dựa trên category_ids/major đã học và weight theo độ hoàn thành.",
            "Popularity smoothing: enrollment_count và rating trung bình làm bias nhẹ để tránh cold item hoàn toàn không xuất hiện.",
            "Learning path booster: nếu user đã hoàn một khóa trong path, khuyến nghị bước k+1 của learning_paths.",
            "Rule layer: không gợi ý khóa đã full complete trừ khi user muốn “ôn lại”; ưu tiên published; filter trạng thái truy cập (free/paid/enrolled).",
        ],
    )

    add_para(doc, "Mô hình ML (giai đoạn nâng cao – tùy dữ liệu)", bold=True)
    add_bullets(
        doc,
        [
            "Matrix factorization / implicit ALS trên Ma trận user×course (từ enroll + partial progress).",
            "Two-tower neural với text + tabular features cho cold-start tốt hơn.",
            "Multi-objective: vừa tối đa completion dự báo vừa đa dạng chủ đề (re-ranking với deterministic diversity penalty).",
        ],
    )

    add_para(doc, "Đầu ra (Output) có cấu trúc", bold=True)
    add_bullets(
        doc,
        [
            "Danh sách ranked items: [{ course_id, score, rationale_tags[], explanation_vi }] — explanation ngắn (vd. “Bạn đang mạnh nhóm Kỹ năng thế kỷ 21”).",
            "Lộ trình gợi ý: path_id hoặc chuỗi course_ids có thứ tự và điều kiện mở khóa (prerequisite có thể mô phỏng bằng rule).",
            "Confidence / bandit: có thể lưu model_version và experimentation_id để đánh giá A/B offline.",
        ],
    )

    add_para(doc, "Cập nhật khi có sự kiện học mới", bold=True)
    add_bullets(
        doc,
        [
            "Near real-time: sau quiz_submit hoặn lesson_completed, phát internal event vào queue; worker cập nhật feature snapshot user.",
            "Nearline: tái huấn luyện model theo lịch (hàng đêm/tuần) khi không yêu cầu latency thấp.",
            "Cold cache: TTL cache gợi ý ngắn (vd. 5–15 phút) với invalidate theo user_id sau sự kiện học đáng kể.",
        ],
    )

    doc.add_heading("4.5. Dashboard và API thống kê", level=2)

    add_para(
        doc,
        "Thiết kế RESTful (hoặc POST body JSON như các route hiện có của dự án) theo vai trò. "
        "Dưới đây là đặc tả gợi ý; endpoint cụ thể có thể map vào controller/service khi triển khai.",
    )

    add_para(doc, "API báo cáo cho sinh viên / người học", bold=True)
    add_bullets(
        doc,
        [
            "GET/POST my-progress-summary: tổng số khóa đang học, % hoàn thành, deadline tự đặt (nếu có).",
            "GET/POST my-quiz-stats: histogram điểm, tỉ lệ pass theo chủ đề.",
            "GET/POST my-study-calendar: hoạt động học (từ study_events aggregated).",
            "GET recommendations: đầu ra module 4.4 (top-k khóa, top-k path).",
        ],
    )

    add_para(doc, "API báo cáo cho quản trị / giảng viên", bold=True)
    add_bullets(
        doc,
        [
            "GET/POST admin/metrics-overview: đăng ký theo ngày, DAU học viên hoạt động (ít nhất một sự kiện học), top khóa theo enroll.",
            "GET/POST admin/course-funnel: preview → enroll → first lesson → complete (cần chuẩn sự kiện).",
            "GET/POST admin/content-health: quiz quá khó (pass rate < ngưỡng), bài video có drop-off.",
            "GET/POST admin/recommendation-kpi: CTR gợi ý, tỉ lệ enroll sau impression (tracking impression_id).",
        ],
    )

    add_para(doc, "Ghi chú hiển thị chỉ báo", bold=True)
    add_bullets(
        doc,
        [
            "Dashboard người học: tập trung động lực (streak, tiến độ, điểm yếu theo chủ đề), tránh quá tải số liệu thô.",
            "Dashboard quản trị: tập trung vận hành (tăng trưởng, chất lượng nội dung, rủi ro bỏ học), phân quyền chặt.",
        ],
    )

    doc.add_heading("4.6. Thiết kế giao diện", level=2)

    add_para(
        doc,
        "Ứng dụng web (React) với layout chính cho người dùng và khu vực /admin cho quản trị. "
        "Dưới đây liệt kê nhóm màn hình chính cần có để đồng bộ với kiến trúc và dữ liệu đã thiết kế.",
    )

    add_para(doc, "Học bài và catalog", bold=True)
    add_bullets(
        doc,
        [
            "Catalog khóa học (published, có phân trang/tìm kiếm server-side).",
            "Chi tiết khóa học: mô tả, giảng viên, curriculum, đăng ký/mua.",
            "Trang học bài: player video / đọc docs / làm quiz–exam có nộp bài và lưu kết quả.",
            "Khóa học của tôi: enroll, tiếp tục học.",
        ],
    )

    add_para(doc, "Gợi ý và lộ trình cá nhân", bold=True)
    add_bullets(
        doc,
        [
            "Trang/block “Gợi ý cho bạn”: danh sách khóa & lộ trình (learning paths) được xếp hạng và giải thích ngắn.",
            "Theo dõi lộ trình: tiến độ từng bước trong bundle; CTA sang khóa kế tiếp.",
            "Wishlist và đánh giá (nếu bật) làm tín hiệu explicit cho gợi ý.",
        ],
    )

    add_para(doc, "Dashboard năng lực (người học)", bold=True)
    add_bullets(
        doc,
        [
            "Bảng điều khiển cá nhân (/dashboard): widget tiến độ, hoạt động gần đây, chỉ báo quiz, gợi ý nhanh.",
            "Chứng chỉ đã đạt (certificates) nếu triển khai hiển thị.",
        ],
    )

    add_para(doc, "Khu quản trị (admin)", bold=True)
    add_bullets(
        doc,
        [
            "Admin dashboard: thống kê tổng quan hệ thống, truy cập nhanh CRUD người dùng/khóa học.",
            "Quản lý nội dung: khóa học, chương, bài, trạng thái xuất bản.",
            "Cấu hình danh mục, ngành/major và (tùy chọn) chiến dụ cụ gợi ý / ngưỡng chỉ báo.",
        ],
    )

    doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.add_run(
        "Ghi chú tài liệu: báo cáo này mô tả thiết kế hệ thống và ánh xạ với các bảng/đường đi "
        "dữ liệu có trong mã nguồn & migration của dự án LMS hiện tại; phần mô hình ML và một số API "
        "thống kê là đặc tả mục tiêu nhằm hoàn chỉnh luồng phân tích và gợi ý theo các mục 4.3–4.5."
    ).italic = True
    footer.paragraph_format.space_after = Pt(6)

    doc.save(OUT)
    print("Wrote:", OUT)


if __name__ == "__main__":
    main()
