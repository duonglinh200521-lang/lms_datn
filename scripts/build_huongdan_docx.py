# -*- coding: utf-8 -*-
"""Generate Report/HuongdanSudung.docx — one-off build script."""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent.parent / "Report" / "HuongdanSudung.docx"


def add_bullets(doc, items):
    for t in items:
        doc.add_paragraph(t, style="List Bullet")


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    title = doc.add_heading("Hướng dẫn sử dụng — Hệ thống LMS", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subt = doc.add_paragraph(
        "Phiên bản tài liệu: theo mã nguồn DTHang (doan_fe-main + doan_be-main)."
    )
    subt.italic = True
    subt.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    doc.add_heading("1. Giới thiệu", level=1)
    doc.add_paragraph(
        "Đây là nền tảng học tập trực tuyến (LMS) với vai trò khách, học viên và quản trị viên. "
        "Giao diện người dùng chạy trên trình duyệt; dữ liệu và nghiệp vụ do API backend cung cấp."
    )

    doc.add_heading("2. Chuẩn bị môi trường và chạy hệ thống", level=1)
    doc.add_heading("2.1. Yêu cầu", level=2)
    add_bullets(
        doc,
        [
            "Node.js 20+ (LTS khuyến nghị)",
            "PostgreSQL 14+",
            "MinIO (tùy chọn — dùng khi upload thumbnail/video; có thể tắt MINIO_DISABLED)",
        ],
    )
    doc.add_heading("2.2. Backend", level=2)
    add_bullets(
        doc,
        [
            "Thư mục: doan_be-main — cài phụ thuộc: npm install",
            "Cấu hình file .env (DB, JWT, MinIO…) theo README.md dự án",
            "Áp dụng migration và seed: npm run db:migrate, npm run db:seed",
            "Chạy API: npm start — mặc định http://localhost:8096/api",
        ],
    )
    doc.add_heading("2.3. Frontend", level=2)
    add_bullets(
        doc,
        [
            "Thư mục: doan_fe-main — npm install",
            "Tạo .env với VITE_API_URL=http://localhost:8096/api (hoặc URL API thực tế)",
            "Chạy dev: npm run dev — mặc định http://localhost:5173",
        ],
    )
    doc.add_paragraph(
        "Đăng nhập sử dụng cookie HTTP-only từ backend; trình duyệt phải cho phép cookie "
        "giữa hai origin FE/BE (đã cấu hình CORS trên server cho localhost)."
    )

    doc.add_heading("3. Vai trò người dùng", level=1)
    t = doc.add_table(rows=4, cols=3)
    t.style = "Table Grid"
    hdr = ("Vai trò", "Đăng nhập", "Khu vực sử dụng chính")
    for j, h in enumerate(hdr):
        t.rows[0].cells[j].text = h
    rows = [
        ("Khách", "Không", "Trang chủ, giới thiệu, xem chi tiết khóa học"),
        ("Học viên", "/auth/login", "Dashboard, khóa của tôi, wishlist, hồ sơ, học trực tuyến"),
        ("Admin", "/auth/admin/login", "Toàn bộ /admin/* (cần role admin)"),
    ]
    for i, (a, b, c) in enumerate(rows, start=1):
        t.rows[i].cells[0].text = a
        t.rows[i].cells[1].text = b
        t.rows[i].cells[2].text = c

    doc.add_heading("4. Hướng dẫn cho khách (chưa đăng nhập)", level=1)
    doc.add_heading("4.1. Trang chủ", level=2)
    doc.add_paragraph("URL: http://localhost:5173/")
    add_bullets(
        doc,
        [
            "Xem giới thiệu, khóa nổi bật và điều hướng tới chi tiết khóa.",
            "Đăng nhập / Đăng ký từ nút trên header khi cần ghi danh hoặc vào khu học viên.",
        ],
    )
    doc.add_heading("4.2. Chi tiết khóa học", level=2)
    doc.add_paragraph("URL: http://localhost:5173/courses/<mã_khóa>")
    add_bullets(
        doc,
        [
            "Đọc mô tả, thông tin khóa và giảng viên.",
            "Sau khi đăng nhập: có thể ghi danh (enroll) và thêm/yêu thích wishlist tùy chức năng hiển thị.",
        ],
    )
    doc.add_heading("4.3. Giới thiệu", level=2)
    doc.add_paragraph("URL: /about — Thông tin giới thiệu nền tảng/khoa.")

    doc.add_heading("5. Hướng dẫn cho học viên", level=1)
    doc.add_heading("5.1. Đăng nhập", level=2)
    doc.add_paragraph("URL: /auth/login — Nhập email và mật khẩu. Sau đăng nhập thành công, "
                      "hệ thống có thể đưa bạn về trang trước đó (nếu có) hoặc tiếp tục thao tác trên site.")
    doc.add_heading("5.2. Bảng điều khiển (Dashboard)", level=2)
    doc.add_paragraph("URL: /dashboard")
    add_bullets(
        doc,
        [
            "Tổng quan: số khóa đã đăng ký, tiến độ, wishlist, gợi ý khóa nổi bật.",
            "Lối tắt đến khóa của tôi, hồ sơ, trang chủ.",
        ],
    )
    doc.add_heading("5.3. Khóa học của tôi", level=2)
    doc.add_paragraph("URL: /my-courses")
    add_bullets(
        doc,
        [
            "Danh sách các khóa đã ghi danh, lọc đang học / hoàn thành, tìm kiếm.",
            "Nút “Tiếp tục học” / “Xem lại” đưa tới màn hình học với tham số courseId.",
        ],
    )
    doc.add_heading("5.4. Wishlist (Yêu thích)", level=2)
    doc.add_paragraph("URL: /my-courses/wishlist — Danh sách khóa đã lưu, mở chi tiết khóa từ danh sách.")
    doc.add_heading("5.5. Học trực tuyến", level=2)
    doc.add_paragraph(
        "URL: /learning?courseId=<id_khóa> — Cần đã đăng nhập và thường phải đã ghi danh khóa."
    )
    add_bullets(
        doc,
        [
            "Thanh bên: danh sách phần và bài học.",
            "Nội dung bài: video, tài liệu (doc), kiểm tra quiz/exam tùy cấu hình.",
            "Tiến độ học và kết quả quiz được lưu về server.",
        ],
    )
    doc.add_heading("5.6. Hồ sơ cá nhân", level=2)
    doc.add_paragraph("URL: /profile — Xem/cập nhật thông tin tài khoản theo chức năng form trên màn hình.")

    doc.add_heading("6. Hướng dẫn cho quản trị viên (Admin)", level=1)
    doc.add_paragraph(
        "Chỉ tài khoản có role admin mới vào được /admin. Đăng nhập tại /auth/admin/login. "
        "Nếu dùng tài khoản không phải admin, hệ thống sẽ báo lỗi quyền và không giữ phiên admin."
    )
    doc.add_heading("6.1. Tổng quan admin", level=2)
    doc.add_paragraph("URL: /admin/dashboard — Thống kê nhanh, đường tắt tới các module, danh sách khóa cập nhật gần đây.")
    doc.add_heading("6.2. Quản lý khóa học", level=2)
    doc.add_paragraph("URL: /admin/v1/courses — Tạo mới, chỉnh sửa, ảnh bìa, trạng thái xuất bản/bản nháp, gán danh mục và chuyên ngành.")
    doc.add_heading("6.3. Nội dung khóa (curriculum)", level=2)
    add_bullets(
        doc,
        [
            "Vào /admin/v1/lesson-management để chọn khóa, sau đó mở chỉnh sửa nội dung.",
            "URL chi tiết: /admin/v1/courses/<courseId>/contents — Thêm/chỉnh section, bài học (video, tài liệu, quiz).",
            "Đồng bộ curriculum về máy chủ sau khi chỉnh sửa.",
        ],
    )
    doc.add_heading("6.4. Người dùng, vai trò, chuyên ngành, danh mục", level=2)
    add_bullets(
        doc,
        [
            "/admin/v1/user-management — Người dùng.",
            "/admin/v1/role-management — Vai trò.",
            "/admin/v1/major-management — Chuyên ngành.",
            "/admin/v1/category-management — Danh mục khóa.",
        ],
    )
    doc.add_heading("6.5. Hồ sơ admin", level=2)
    doc.add_paragraph("URL: /admin/v1/profile")

    doc.add_heading("7. Chuyển hướng URL thường gặp", level=1)
    add_bullets(
        doc,
        [
            "/admin → chuyển tới /admin/dashboard",
            "/admin/v1 → chuyển tới /admin/v1/courses",
        ],
    )

    doc.add_heading("8. Xử lý sự cố thường gặp", level=1)
    add_bullets(
        doc,
        [
            "Đăng nhập xong không gọi được API: kiểm tra VITE_API_URL, backend có chạy, CORS và cookie (withCredentials).",
            "Không vào được trang học viên/admin: đảm bảo đúng luồng đăng nhập (/auth/login vs /auth/admin/login).",
            "Upload ảnh/video lỗi: kiểm tra MinIO và biến môi trường hoặc tạm đặt MINIO_DISABLED và tránh chức năng upload.",
            "Lỗi kết nối Postgres: kiểm tra .env và quyền user CSDL; chạy lại migration.",
        ],
    )

    doc.add_heading("9. Tài liệu tham chiếu trong mã nguồn", level=1)
    add_bullets(
        doc,
        [
            "README.md — Cài đặt, biến môi trường, migration, build.",
            "Function.md — Bảng màn hình và nhóm API.",
            "doan_fe-main/src/routes/index.tsx — Định nghĩa route và bảo vệ quyền.",
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph("_Hết — Mở file này bằng Microsoft Word hoặc WPS để chỉnh sửa định dạng báo cáo._")
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(10)

    doc.save(str(OUT))
    print("Saved:", OUT)


if __name__ == "__main__":
    main()
