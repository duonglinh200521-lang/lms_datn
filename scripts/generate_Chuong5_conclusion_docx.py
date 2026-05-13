# -*- coding: utf-8 -*-
"""Tạo CHƯƠNG 5: Kết luận và định hướng (Chuong5.docx)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent / "Chuong5.docx"


def bullets(doc, items):
    for t in items:
        p = doc.add_paragraph(t, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        for r in p.runs:
            r.font.size = Pt(12)


def main():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Pt(72)

    h = doc.add_heading(
        "CHƯƠNG 5: KẾT LUẬN VÀ ĐỊNH HƯỚNG PHÁT TRIỂN", level=1
    )
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 5.1
    doc.add_heading("5.1. Kết quả đạt được", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Trong phạm vi đồ án xây dựng hệ thống LMS hướng nghiệp (học trực tuyến, quản lý khóa học và "
        "người dùng), nhóm thực hiện đã hoàn thành một chuỗi sản phẩm có thể vận hành thử nghiệm, "
        "gồm backend REST API, cơ sở dữ liệu quan hệ và giao diện web SPA. Các kết quả chính có thể tóm tắt như sau:"
    )
    for r in p.runs:
        r.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(8)

    bullets(
        doc,
        [
            "Hoàn thiện kiến trúc Client–Server: frontend React (Vite, Ant Design) giao tiếp với backend Node.js/Express/TypeScript; phân tách rõ tầng xử lý nghiệp vụ, repository và lưu trữ PostgreSQL.",
            "Xây dựng luồng nghiệp vụ lõi LMS: quản lý khóa học, chương–bài học (video, tài liệu, trắc nghiệm, bài kiểm tra), trạng thái xuất bản, đăng ký học và phân quyền theo vai trò (admin, giảng viên, học viên) với xác thực JWT/cookie.",
            "Bổ sung danh mục công khai có phân trang: API catalog cho khóa học đã xuất bản, hỗ trợ tìm kiếm và lọc theo ngành/danh mục, phục vụ trải nghiệm người dùng chưa đăng nhập hoặc khám phá nội dung.",
            "Thiết kế và triển khai schema dữ liệu phục vụ theo dõi học tập: tiến độ theo bài (lesson_progress), lịch sử làm quiz/exam (user_quiz_attempts), khả năng ghi nhận sự kiện học (study_events) làm nền cho phân tích và gợi ý sau này.",
            "Mở rộng nội dung minh họa: seed dữ liệu demo gồm nhiều nhóm khóa hướng nghiệp, kỹ năng và lộ trình combo; hỗ trợ minh họa báo cáo, demo giao diện catalog và quản trị.",
            "Tích hợp tùy chọn lưu trữ media (MinIO) cho ảnh/video; tài liệu hóa quy trình cài đặt, migration và seed trong kho mã (README, script migrate/seed).",
        ],
    )

    doc.add_paragraph()

    # 5.2
    doc.add_heading("5.2. Hạn chế của đề tài", level=2)
    p2 = doc.add_paragraph()
    p2.add_run(
        "Bên cạnh các kết quả đạt được, đồ án vẫn tồn tại những hạn chế khách quan và chủ quan, cần được nêu rõ để làm cơ sở cho các hướng hoàn thiện tiếp theo:"
    )
    for r in p2.runs:
        r.font.size = Pt(12)
    p2.paragraph_format.space_after = Pt(8)

    bullets(
        doc,
        [
            "Module gợi ý thông minh (AI Recommendation) mới dừng ở mức thiết kế hoặc heuristic đơn giản: chưa có mô hình học máy đầy đủ được huấn luyện và đánh giá trên dữ liệu thực tế người dùng đủ lớn.",
            "Thu thập hành vi học chi tiết (heartbeat video, funnel từng bước, A/B impression) chưa được triển khai đồng bộ trên toàn bộ luồng UI—khiến một số chỉ báo trong thiết kế phân tích chưa có đủ dữ liệu thô.",
            "Dashboard và API thống kê đa chiều (học viên, khóa học, vận hành hệ thống) có thể cần thêm điểm cuối tổng hợp, cache và tối ưu báo cáo khi khối lượng người dùng và sự kiện tăng.",
            "Bài toán thanh toán, chứng chỉ, phản hồi đánh giá khóa học và các tình huống biên (nhiều thiết bị, gián đoạn mạng) cần kiểm thử sâu và bổ sung kịch bản nghiệp vụ trong môi trường production.",
            "Tài liệu kiểm thử tự động (unit/integration/E2E) và pipeline CI/CD chưa được nhấn mạnh trong phạm vi đồ án hiện tại.",
            "Độ bao phủ bảo mật (quan trắc, pentest định kỳ, chính sách GDPR nếu triển khai quốc tế) cần được nâng cấp khi đưa hệ thống ra sử dụng rộng rãi.",
        ],
    )

    doc.add_paragraph()

    # 5.3
    doc.add_heading("5.3. Hướng phát triển của đề tài", level=2)
    p3 = doc.add_paragraph()
    p3.add_run(
        "Xuất phát từ kết quả và hạn chế nêu trên, định hướng phát triển có thể được sắp xếp theo thứ tự ưu tiên kỹ thuật–sản phẩm như sau:"
    )
    for r in p3.runs:
        r.font.size = Pt(12)
    p3.paragraph_format.space_after = Pt(8)

    bullets(
        doc,
        [
            "Hoàn thiện pipeline dữ liệu: chuẩn hóa schema payload cho study_events, job tổng hợp chỉ báo định kỳ, materialized view hoặc bảng feature phục vụ gợi ý và dashboard.",
            "Triển khai dịch vụ gợi ý theo lộ trình: giai đoạn 1 content-based + popularity + lộ trình learning paths; giai đoạn 2 mô hình implicit feedback (ALS, two-tower) khi đủ log enroll và tiến độ.",
            "Mở rộng API và giao diện dashboard: biểu đồ cho học viên (streak, điểm yếu theo chủ đề) và cho quản trị (DAU học, funnel nội dung, chất lượng quiz).",
            "Tăng cường trải nghiệm học: offline-friendly cho tài liệu, cảnh báo khi mất kết nối, thông báo nhắc học (email/push) tùy cấu hình.",
            "Chuẩn hóa triển khai: Docker Compose hoặc Kubernetes cho production, reverse proxy (HTTPS), backup DB định kỳ, giám sát log và metrics (Prometheus/Grafana hoặc tương đương).",
            "Mở rộng tích hợp: SSO doanh nghiệp/giáo dục, API webhook cho hệ thống bên ngoài, xuất báo cáo PDF/Excel cho phụ huynh hoặc cố vấn học tập.",
        ],
    )

    doc.add_paragraph()
    fin = doc.add_paragraph()
    fin.add_run(
        "Kết luận chung: đồ án đã chứng minh khả năng xây dựng một LMS hiện đại với quy trình phát triển rõ ràng và nền tảng dữ liệu phù hợp mở rộng. "
        "Các hướng phát triển trên giúp hệ thống tiệm cận sản phẩm thương mại hoặc nghiên cứu ứng dụng AI trong giáo dục định hướng nghề nghiệp."
    ).italic = True
    for r in fin.runs:
        r.font.size = Pt(12)

    doc.save(OUT)
    print("Wrote:", OUT)


if __name__ == "__main__":
    main()
